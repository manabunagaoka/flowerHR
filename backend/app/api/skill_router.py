import io
import json
from typing import Optional, List
from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import PlainTextResponse, StreamingResponse
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.services.task_enricher import parse_jd_to_template, generate_skill_md_from_template, polish_text
from app.services.doc_parser import parse_document

router = APIRouter()


class JDInput(BaseModel):
    job_title: str
    job_description: str
    daily_tasks: List[str]


class TaskTemplate(BaseModel):
    id: int
    name: str
    description: str = ""
    frequency: str = ""
    tools_systems: str = ""
    file_paths: str = ""
    contacts: str = ""
    kpi: str = ""
    automation_potential: str = "unknown"


class ProjectTemplate(BaseModel):
    id: int
    name: str
    description: str = ""
    tasks: List[TaskTemplate]


class TemplateInput(BaseModel):
    job_title: str
    department: str = ""
    summary: str = ""
    employee_name: Optional[str] = None
    projects: List[ProjectTemplate]
    qualifications: List[str] = []
    skills_mentioned: List[str] = []


@router.post("/parse-jd")
async def parse_jd(input: JDInput):
    try:
        template = await parse_jd_to_template(
            job_title=input.job_title,
            job_description=input.job_description,
            daily_tasks=input.daily_tasks,
        )
        return template
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/upload-jd")
async def upload_jd(file: UploadFile = File(...)):
    if not file.filename.endswith((".docx", ".pdf")):
        raise HTTPException(status_code=400, detail="Only .docx and .pdf files are supported")

    try:
        content = await file.read()
        parsed = parse_document(content, file.filename)

        template = await parse_jd_to_template(
            job_title=parsed["job_title"],
            job_description=parsed["job_description"],
            daily_tasks=parsed["daily_tasks"],
        )
        return template
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


class PolishInput(BaseModel):
    text: str


@router.post("/polish", response_class=PlainTextResponse)
async def polish(input: PolishInput):
    try:
        return await polish_text(input.text)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/generate-skill", response_class=PlainTextResponse)
async def generate_skill(input: TemplateInput):
    try:
        skill_md = await generate_skill_md_from_template(
            template=input.model_dump(),
            employee_name=input.employee_name,
        )
        return skill_md
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/download-jd-docx")
async def download_jd_docx(input: TemplateInput):
    try:
        doc = Document()

        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        title = doc.add_heading("JOB DESCRIPTION", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("")
        info = doc.add_paragraph()
        info.add_run("Employee: ").bold = True
        info.add_run(input.employee_name or "TBD")
        info = doc.add_paragraph()
        info.add_run("Job Title: ").bold = True
        info.add_run(input.job_title)
        info = doc.add_paragraph()
        info.add_run("Department: ").bold = True
        info.add_run(input.department)

        doc.add_heading("Summary", level=1)
        doc.add_paragraph(input.summary)

        for project in input.projects:
            doc.add_heading(project.name, level=1)
            if project.description:
                desc = doc.add_paragraph(project.description)
                desc.style = doc.styles["Normal"]

            for task in project.tasks:
                doc.add_heading(task.name, level=2)

                if task.description:
                    doc.add_paragraph(task.description)

                details = []
                if task.frequency:
                    details.append(("Frequency", task.frequency))
                if task.tools_systems:
                    details.append(("Tools/Systems", task.tools_systems))
                if task.file_paths:
                    details.append(("File Locations", task.file_paths))
                if task.contacts:
                    details.append(("Key Contacts", task.contacts))
                if task.kpi:
                    details.append(("KPI", task.kpi))

                for label, value in details:
                    p = doc.add_paragraph()
                    p.add_run(f"{label}: ").bold = True
                    p.add_run(value)

        if input.qualifications:
            doc.add_heading("Qualifications", level=1)
            for q in input.qualifications:
                doc.add_paragraph(q, style="List Bullet")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"JD_{(input.employee_name or input.job_title).replace(' ', '_')}.docx"
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/generate-task-template")
async def generate_task_template(input: TemplateInput):
    try:
        doc = Document()

        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        title = doc.add_heading("PROJECTS & TASKS TEMPLATE", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("")
        info = doc.add_paragraph()
        info.add_run("Employee: ").bold = True
        info.add_run(input.employee_name or "___________________________")

        info = doc.add_paragraph()
        info.add_run(f"Generated from JD: {input.job_title}")

        doc.add_paragraph("")

        instructions = doc.add_paragraph()
        run = instructions.add_run(
            "Instructions: For each project area below, describe the tasks you perform. "
            "Fill in as much detail as you can. Add extra tasks if needed."
        )
        run.italic = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph("")

        for project in input.projects:
            doc.add_heading(f"PROJECT: {project.name}", level=1)

            if project.description:
                desc_para = doc.add_paragraph(project.description)
                desc_para.style = doc.styles["Normal"]

            hint = doc.add_paragraph()
            run = hint.add_run("(describe the tasks you do under this area)")
            run.italic = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

            for task_num in range(1, 4):
                doc.add_heading(f"Task {task_num}", level=2)

                fields = [
                    ("Task Name", ""),
                    ("How often (Frequency)", ""),
                    ("Tools/Systems used", ""),
                    ("Key contacts", ""),
                    ("How success is measured (KPI)", ""),
                ]
                for label, value in fields:
                    p = doc.add_paragraph()
                    p.add_run(f"{label}: ").bold = True
                    p.add_run(value)

                doc.add_paragraph("")

        doc.add_paragraph("")
        doc.add_heading("Additional Projects", level=1)
        additional = doc.add_paragraph()
        run = additional.add_run(
            "Add more projects below if your work includes areas not covered in the JD above."
        )
        run.italic = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph("")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"Task_Template_{(input.employee_name or input.job_title).replace(' ', '_')}.docx"
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/parse-task-doc")
async def parse_task_doc(file: UploadFile = File(...)):
    if not file.filename.endswith((".docx", ".pdf")):
        raise HTTPException(status_code=400, detail="Only .docx and .pdf files are supported")

    try:
        content = await file.read()
        parsed = parse_document(content, file.filename)

        template = await parse_jd_to_template(
            job_title=parsed["job_title"],
            job_description=parsed["job_description"],
            daily_tasks=parsed["daily_tasks"],
        )

        return template
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/download-tasks-docx")
async def download_tasks_docx(input: TemplateInput):
    try:
        doc = Document()

        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        title = doc.add_heading("PROJECTS & TASKS", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("")
        info = doc.add_paragraph()
        info.add_run("Employee: ").bold = True
        info.add_run(input.employee_name or "TBD")
        info = doc.add_paragraph()
        info.add_run("Role: ").bold = True
        info.add_run(input.job_title)
        info = doc.add_paragraph()
        info.add_run("Department: ").bold = True
        info.add_run(input.department)

        doc.add_paragraph("")

        for project in input.projects:
            doc.add_heading(project.name, level=1)
            if project.description:
                desc = doc.add_paragraph(project.description)
                desc.style = doc.styles["Normal"]

            for task in project.tasks:
                doc.add_heading(task.name, level=2)

                if task.description:
                    doc.add_paragraph(task.description)

                details = [
                    ("Frequency", task.frequency),
                    ("Tools/Systems", task.tools_systems),
                    ("File Locations", task.file_paths),
                    ("Key Contacts", task.contacts),
                    ("KPI", task.kpi),
                    ("Automation Potential", task.automation_potential),
                ]

                for label, value in details:
                    if value and value != "unknown":
                        p = doc.add_paragraph()
                        p.add_run(f"{label}: ").bold = True
                        p.add_run(value)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"Tasks_{(input.employee_name or input.job_title).replace(' ', '_')}.docx"
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/download-tasks-json")
async def download_tasks_json(input: TemplateInput):
    try:
        data = {
            "employee": input.employee_name or "",
            "role": input.job_title,
            "department": input.department,
            "projects": [
                {
                    "name": project.name,
                    "description": project.description,
                    "tasks": [
                        {
                            "name": task.name,
                            "description": task.description,
                            "frequency": task.frequency,
                            "tools_systems": task.tools_systems,
                            "file_paths": task.file_paths,
                            "contacts": task.contacts,
                            "kpi": task.kpi,
                            "automation_potential": task.automation_potential,
                        }
                        for task in project.tasks
                    ],
                }
                for project in input.projects
            ],
        }

        json_bytes = json.dumps(data, indent=2, ensure_ascii=False).encode("utf-8")
        buffer = io.BytesIO(json_bytes)
        buffer.seek(0)

        filename = f"Tasks_{(input.employee_name or input.job_title).replace(' ', '_')}.json"
        return StreamingResponse(
            buffer,
            media_type="application/json",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
