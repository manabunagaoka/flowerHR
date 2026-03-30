import io
import re
from docx import Document
import pdfplumber


def parse_document(content: bytes, filename: str) -> dict:
    if filename.endswith(".docx"):
        return parse_docx(content)
    elif filename.endswith(".pdf"):
        return parse_pdf(content)
    raise ValueError(f"Unsupported file type: {filename}")


def parse_docx(content: bytes) -> dict:
    doc = Document(io.BytesIO(content))

    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text.strip())

    # Also extract from tables (common in JD templates)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    full_text.append(text)

    return extract_sections("\n".join(full_text))


def parse_pdf(content: bytes) -> dict:
    full_text = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                full_text.append(text)

    return extract_sections("\n".join(full_text))


def extract_sections(text: str) -> dict:
    """Extract job title, description, and tasks from raw document text.

    Uses common JD section headers to find structure. Falls back to
    sending the full text as the description if no clear sections found.
    """
    lines = [l.strip() for l in text.split("\n") if l.strip()]

    job_title = ""
    job_description_lines = []
    task_lines = []

    # Common header patterns
    title_patterns = re.compile(
        r"^(job\s*title|position\s*title|role|title)\s*[:\-]?\s*(.+)",
        re.IGNORECASE,
    )
    task_section_patterns = re.compile(
        r"^(key\s*responsibilities|responsibilities|duties|daily\s*tasks|tasks|day[- ]to[- ]day)",
        re.IGNORECASE,
    )
    desc_section_patterns = re.compile(
        r"^(job\s*description|description|summary|overview|about\s*the\s*role|purpose)",
        re.IGNORECASE,
    )

    current_section = None

    for line in lines:
        title_match = title_patterns.match(line)
        if title_match:
            job_title = title_match.group(2).strip()
            continue

        if task_section_patterns.match(line):
            current_section = "tasks"
            continue

        if desc_section_patterns.match(line):
            current_section = "description"
            continue

        if current_section == "tasks":
            cleaned = re.sub(r"^[\-\u2022\*\d\.]+\s*", "", line)
            if cleaned:
                task_lines.append(cleaned)
        elif current_section == "description":
            job_description_lines.append(line)
        elif not job_title and not current_section:
            # First substantial line might be the title
            if len(line) < 80 and not job_title:
                job_title = line

    # Fallback: if no structured sections found, use full text as description
    if not job_description_lines:
        job_description_lines = lines

    return {
        "job_title": job_title or "Unknown Role",
        "job_description": "\n".join(job_description_lines),
        "daily_tasks": task_lines if task_lines else ["(No tasks extracted — please review document)"],
    }
